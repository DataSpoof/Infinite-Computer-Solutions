# Install required libraries (Run this only once)
# pip install openpyxl python-docx
# uv add openpyxl python-docx

# winrar
# termux
# sublime
# ipynb viewer
# chrome remote desktop


#Excel-- openpyxl
#CSV, parquet, tsv, excel--- pandas (small data)
# LARGE DATASET (POLARS, PYSPARK, DASK)
# PDF-- (pymupdf)
# AUDIO--- LIBROSA
# image-- skimage, cv2
# pydicom-- reading x ray images in dcm format
# rasterio--  satelite kind of images

# extracting the table from a pdf we use---- microsoft table transformer
#-- bordered table 
#-- borderless tables
#-- key value











# Import Workbook class to create a new Excel workbook
# Import load_workbook to open an existing workbook
from openpyxl import Workbook, load_workbook

# Import Font for changing text style
# Import PatternFill for adding background color to cells
from openpyxl.styles import Font, PatternFill

# Import Document class to create Word documents
from docx import Document

# Import Pt to specify font size in points (not used in this example)
from docx.shared import Pt


# =====================================================
# STEP 1 : CREATE AN EXCEL WORKBOOK
# =====================================================

# Create a new Excel workbook
wb = Workbook()

# Select the default worksheet
emp_sheet = wb.active

# Rename the default worksheet
emp_sheet.title = "Employees"

# Add header row
emp_sheet.append(["ID", "Name", "Department", "Salary"])

# Employee data stored as a list of lists
employees = [
    [101, "Alice", "IT", 70000],
    [102, "Bob", "HR", 55000],
    [103, "Charlie", "Finance", 65000],
    [104, "David", "IT", 90000],
    [105, "Eva", "Sales", 60000],
]

# Insert each employee into Excel
for row in employees:
    emp_sheet.append(row)

# Apply formatting to header row
for cell in emp_sheet[1]:

    # Make header text bold and white
    cell.font = Font(
        bold=True,
        color="FFFFFF"
    )

    # Fill header with blue background
    cell.fill = PatternFill(
        fill_type="solid",
        fgColor="4F81BD"
    )


# =====================================================
# STEP 2 : CREATE SUMMARY WORKSHEET
# =====================================================

# Create another worksheet
summary = wb.create_sheet("Summary")

# Add column headings
summary["A1"] = "Department"
summary["B1"] = "Total Salary"

# Make headings bold
summary["A1"].font = Font(bold=True)
summary["B1"].font = Font(bold=True)

# Dictionary to store department-wise salary totals
department_salary = {}

# Calculate salary totals
for row in employees:

    # Extract department
    dept = row[2]

    # Extract salary
    salary = row[3]

    # Add salary to department total
    department_salary[dept] = (
        department_salary.get(dept, 0) + salary
    )

# Start writing from second row
row_num = 2

# Write summary into worksheet
for dept, total in department_salary.items():

    summary.cell(
        row=row_num,
        column=1
    ).value = dept

    summary.cell(
        row=row_num,
        column=2
    ).value = total

    row_num += 1

# Excel filename
excel_file = "company_report.xlsx"

# Save workbook
wb.save(excel_file)

print("Excel workbook created successfully!")











# =====================================================
# STEP 3 : READ THE EXCEL FILE
# =====================================================

# Load existing workbook
wb = load_workbook(excel_file)

# Open Employees worksheet
employees_sheet = wb["Employees"]

# Open Summary worksheet
summary_sheet = wb["Summary"]

print("\nEmployees:\n")

# Read all employee records
for row in employees_sheet.iter_rows(
    min_row=2,
    values_only=True
):
    print(row)


# =====================================================
# STEP 4 : CREATE WORD REPORT
# =====================================================

# Create Word document
document = Document()

# Main heading
document.add_heading(
    "Company Employee Report",
    level=1
)

# Section heading
document.add_heading(
    "Employee Details",
    level=2
)

# Create table with one header row and four columns
table = document.add_table(
    rows=1,
    cols=4
)

# Apply table style
table.style = "Table Grid"

# Access header cells
hdr = table.rows[0].cells

# Add header text
hdr[0].text = "ID"
hdr[1].text = "Name"
hdr[2].text = "Department"
hdr[3].text = "Salary"

# Add employee data into Word table
for row in employees_sheet.iter_rows(
    min_row=2,
    values_only=True
):

    # Create new row
    cells = table.add_row().cells

    # Populate cells
    cells[0].text = str(row[0])
    cells[1].text = row[1]
    cells[2].text = row[2]
    cells[3].text = str(row[3])

# Insert page break
document.add_page_break()

# Add heading
document.add_heading(
    "Department Salary Summary",
    level=2
)

# Create summary table
summary_table = document.add_table(
    rows=1,
    cols=2
)

summary_table.style = "Light Grid"

# Header row
header = summary_table.rows[0].cells

header[0].text = "Department"
header[1].text = "Total Salary"

# Add department summary
for row in summary_sheet.iter_rows(
    min_row=2,
    values_only=True
):

    cells = summary_table.add_row().cells

    cells[0].text = row[0]
    cells[1].text = str(row[1])

# Insert another page break
document.add_page_break()

# Add statistics heading
document.add_heading(
    "Statistics",
    level=2
)

# Initialize variables
total_salary = 0
employee_count = 0

# Calculate statistics
for row in employees_sheet.iter_rows(
    min_row=2,
    values_only=True
):

    employee_count += 1
    total_salary += row[3]

# Compute average salary
average_salary = total_salary / employee_count

# Create paragraph
p = document.add_paragraph()

# Total employees
p.add_run("Total Employees : ").bold = True
p.add_run(str(employee_count))

# Total salary
p.add_run("\nTotal Salary : ").bold = True
p.add_run(str(total_salary))

# Average salary
p.add_run("\nAverage Salary : ").bold = True
p.add_run(f"{average_salary:.2f}")

# Output Word filename
word_file = "Employee_Report.docx"

# Save document
document.save(word_file)

print("\nWord report generated successfully!")